"""
Crée les fichiers d'exemple :
  - config/questions.xlsx   (questions de démonstration)
  - templates/template.docx          (template standard)
  - templates/template_mineur.docx   (variante mineur)

Lancez ce script une seule fois :
    python setup_sample_data.py
"""

import os

import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Helpers Excel
# ---------------------------------------------------------------------------

def _header_style(cell, bg_hex: str = "1a3a5c") -> None:
    cell.font      = Font(bold=True, color="FFFFFF", size=11)
    cell.fill      = PatternFill("solid", fgColor=bg_hex)
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def _auto_width(ws) -> None:
    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 60)


# ---------------------------------------------------------------------------
# Données des questions (exemple : assignation en référé)
# ---------------------------------------------------------------------------

QUESTIONS = [
    # id, question, type, options, variable, show_if, section, required
    # --- DEMANDEUR ---
    ("Q001", "Nom et prénom du demandeur", "text",  "", "nom_demandeur",  "", "Demandeur", "oui"),
    ("Q002", "Adresse du demandeur",        "multiline", "", "adresse_demandeur", "", "Demandeur", "oui"),
    ("Q003", "Qualité du demandeur",        "choice",
     "Particulier|Société|Association|Établissement public",
     "qualite_demandeur", "", "Demandeur", "oui"),
    ("Q004", "SIRET (si société)",          "text",  "", "siret_demandeur",
     "Q003=Société", "Demandeur", "non"),
    ("Q005", "Représenté par (avocat / mandataire)", "text", "", "representant_demandeur",
     "", "Demandeur", "non"),

    # --- DÉFENDEUR ---
    ("Q006", "Nom et prénom / raison sociale du défendeur", "text", "", "nom_defendeur",
     "", "Défendeur", "oui"),
    ("Q007", "Adresse du défendeur",        "multiline", "", "adresse_defendeur",
     "", "Défendeur", "oui"),
    ("Q008", "Qualité du défendeur",        "choice",
     "Particulier|Société|Association|Établissement public",
     "qualite_defendeur", "", "Défendeur", "oui"),
    ("Q009", "SIRET du défendeur (si société)", "text", "", "siret_defendeur",
     "Q008=Société", "Défendeur", "non"),
    ("Q010", "Le défendeur est-il mineur ?","yes_no", "", "est_mineur",
     "", "Défendeur", "oui"),
    ("Q011", "Nom du représentant légal (si mineur)", "text", "", "representant_legal",
     "Q010=oui", "Défendeur", "oui"),
    ("Q012", "Adresse du représentant légal", "multiline", "", "adresse_representant_legal",
     "Q010=oui", "Défendeur", "oui"),

    # --- JURIDICTION ---
    ("Q013", "Juridiction saisie",          "choice",
     "Tribunal judiciaire|Tribunal de commerce|Conseil de prud'hommes|Cour d'appel",
     "juridiction", "", "Juridiction", "oui"),
    ("Q014", "Ville du tribunal",           "text",  "", "ville_tribunal",
     "", "Juridiction", "oui"),
    ("Q015", "Date d'audience souhaitée",   "date",  "", "date_audience",
     "", "Juridiction", "oui"),
    ("Q016", "Heure d'audience",            "text",  "", "heure_audience",
     "", "Juridiction", "oui"),
    ("Q017", "Chambre / service",           "text",  "", "chambre",
     "", "Juridiction", "non"),

    # --- OBJET DU LITIGE ---
    ("Q018", "Nature du litige",            "choice",
     "Impayé|Expulsion|Travaux non réalisés|Trouble de voisinage|Rupture de contrat|Autre",
     "nature_litige", "", "Litige", "oui"),
    ("Q019", "Précisez (si Autre)",         "text",  "", "nature_litige_autre",
     "Q018=Autre", "Litige", "oui"),
    ("Q020", "Date de naissance du litige (contrat, événement…)", "date", "", "date_litige",
     "", "Litige", "oui"),
    ("Q021", "Description des faits",       "multiline", "", "description_faits",
     "", "Litige", "oui"),
    ("Q022", "Montant réclamé (en euros)",  "number", "", "montant_reclame",
     "", "Litige", "non"),
    ("Q023", "Base juridique principale",   "text",  "", "base_juridique",
     "", "Litige", "non"),

    # --- PIÈCES ---
    ("Q024", "Liste des pièces jointes",    "multiline", "", "pieces_jointes",
     "", "Pièces", "non"),

    # --- DEMANDES ---
    ("Q025", "Demandes au fond",            "multiline", "", "demandes_fond",
     "", "Demandes", "oui"),
    ("Q026", "Demande de provision (oui/non)", "yes_no", "", "demande_provision",
     "", "Demandes", "oui"),
    ("Q027", "Montant de la provision demandée", "number", "", "montant_provision",
     "Q026=oui", "Demandes", "oui"),
    ("Q028", "Demande d'article 700 CPC (oui/non)", "yes_no", "", "demande_700",
     "", "Demandes", "oui"),
    ("Q029", "Montant article 700 demandé", "number", "", "montant_700",
     "Q028=oui", "Demandes", "oui"),
    ("Q030", "Demande d'exécution provisoire", "yes_no", "", "execution_provisoire",
     "", "Demandes", "non"),

    # --- SIGNATURE ---
    ("Q031", "Ville de rédaction",          "text",  "", "ville_redaction",
     "", "Signature", "oui"),
    ("Q032", "Date de rédaction",           "date",  "", "date_redaction",
     "", "Signature", "oui"),
    ("Q033", "Nom du rédacteur",            "text",  "", "nom_redacteur",
     "", "Signature", "oui"),
]

VARIANTS = [
    # condition, template_file
    ("Q010=oui", "template_mineur.docx"),
    ("Q013=Tribunal de commerce", "template_commerce.docx"),
]


# ---------------------------------------------------------------------------
# Création du fichier Excel
# ---------------------------------------------------------------------------

def create_questions_excel(path: str) -> None:
    wb = openpyxl.Workbook()

    # --- Feuille Questions ---
    ws = wb.active
    ws.title = "Questions"
    ws.row_dimensions[1].height = 30

    columns = ["id", "question", "type", "options", "variable",
               "show_if", "section", "required"]
    ws.append(columns)
    for i, col in enumerate(columns, 1):
        _header_style(ws.cell(1, i))

    for q in QUESTIONS:
        ws.append(list(q))

    # Alternance de couleurs
    for row_idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
        fill = PatternFill("solid", fgColor="EBF5FB" if row_idx % 2 == 0 else "FFFFFF")
        for cell in row:
            cell.fill = fill
            cell.alignment = Alignment(wrap_text=True, vertical="center")

    _auto_width(ws)

    # --- Feuille Variants ---
    ws_v = wb.create_sheet("Variants")
    ws_v.row_dimensions[1].height = 30
    ws_v.append(["condition", "template_file"])
    for i in range(1, 3):
        _header_style(ws_v.cell(1, i), "16324f")
    for v in VARIANTS:
        ws_v.append(list(v))
    _auto_width(ws_v)

    # --- Feuille Aide ---
    ws_h = wb.create_sheet("Aide")
    aide = [
        ["Colonne",     "Description"],
        ["id",          "Identifiant unique de la question (ex: Q001)"],
        ["question",    "Texte affiché à l'utilisateur"],
        ["type",        "text | date | number | multiline | yes_no | choice"],
        ["options",     "Pour type=choice : valeurs séparées par |  (ex: Oui|Non|Peut-être)"],
        ["variable",    "Nom de la variable dans le template Word  {{ nom_variable }}"],
        ["show_if",     "Condition d'affichage (ex: Q005=oui  ou  Q005!=non)  — vide = toujours"],
        ["section",     "Nom du groupe/section affiché dans l'interface"],
        ["required",    "oui → réponse obligatoire,  non → peut être passé"],
    ]
    for row in aide:
        ws_h.append(row)
    ws_h["A1"].font = Font(bold=True)
    ws_h["B1"].font = Font(bold=True)
    ws_h.column_dimensions["A"].width = 18
    ws_h.column_dimensions["B"].width = 70

    os.makedirs(os.path.dirname(path), exist_ok=True)
    wb.save(path)
    print(f"  ✓  {path}")


# ---------------------------------------------------------------------------
# Création d'un template Word
# ---------------------------------------------------------------------------

def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)


def create_template(path: str, variant: str = "standard") -> None:
    """
    Crée un template Word avec des variables Jinja2 {{ variable }}.
    Pour un vrai usage, remplacez ce template par votre propre modèle
    et ajoutez les {{ variables }} aux endroits souhaités.
    """
    doc = Document()

    # Style général
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ---- En-tête ----
    if variant == "mineur":
        doc.add_paragraph(
            "ASSIGNATION EN RÉFÉRÉ\n(Défendeur mineur — représenté par son représentant légal)"
        ).runs[0].bold = True
    else:
        doc.add_paragraph("ASSIGNATION").runs[0].bold = True

    doc.add_paragraph()

    # ---- Parties ----
    _add_heading(doc, "PARTIES", 2)

    doc.add_paragraph(
        "DEMANDEUR : {{ nom_demandeur }}{% if qualite_demandeur %} ({{ qualite_demandeur }}){% endif %}"
    )
    doc.add_paragraph("Demeurant : {{ adresse_demandeur }}")
    if variant != "mineur":
        doc.add_paragraph(
            "{% if representant_demandeur %}Représenté par : {{ representant_demandeur }}{% endif %}"
        )

    doc.add_paragraph()

    doc.add_paragraph(
        "DÉFENDEUR : {{ nom_defendeur }}{% if qualite_defendeur %} ({{ qualite_defendeur }}){% endif %}"
    )
    doc.add_paragraph("Demeurant : {{ adresse_defendeur }}")

    if variant == "mineur":
        doc.add_paragraph(
            "Représenté par son représentant légal : {{ representant_legal }}"
        )
        doc.add_paragraph("Demeurant : {{ adresse_representant_legal }}")

    doc.add_paragraph()

    # ---- Juridiction ----
    _add_heading(doc, "JURIDICTION", 2)
    doc.add_paragraph(
        "À Monsieur/Madame le Président du {{ juridiction }} de {{ ville_tribunal }}"
    )
    doc.add_paragraph(
        "Audience du {{ date_audience }} à {{ heure_audience }}"
        "{% if chambre %}, {{ chambre }}{% endif %}"
    )

    doc.add_paragraph()

    # ---- Faits ----
    _add_heading(doc, "EXPOSÉ DES FAITS", 2)
    doc.add_paragraph(
        "Le présent litige, de nature « {{ nature_litige }} », est né le {{ date_litige }}."
    )
    doc.add_paragraph("{{ description_faits }}")

    doc.add_paragraph()

    # ---- Base juridique ----
    _add_heading(doc, "BASE JURIDIQUE", 2)
    doc.add_paragraph(
        "{% if base_juridique %}{{ base_juridique }}{% else %}[À compléter]{% endif %}"
    )

    doc.add_paragraph()

    # ---- Demandes ----
    _add_heading(doc, "DEMANDES", 2)
    doc.add_paragraph("{{ demandes_fond }}")

    doc.add_paragraph(
        "{% if demande_provision == 'oui' %}"
        "Il est demandé une provision de {{ montant_provision }} €."
        "{% endif %}"
    )
    doc.add_paragraph(
        "{% if demande_700 == 'oui' %}"
        "Il est demandé la condamnation aux frais et dépens ainsi qu'à la somme de "
        "{{ montant_700 }} € au titre de l'article 700 du CPC."
        "{% endif %}"
    )
    doc.add_paragraph(
        "{% if montant_reclame %}"
        "Montant total réclamé : {{ montant_reclame }} €."
        "{% endif %}"
    )

    doc.add_paragraph()

    # ---- Pièces ----
    _add_heading(doc, "PIÈCES PRODUITES", 2)
    doc.add_paragraph(
        "{% if pieces_jointes %}{{ pieces_jointes }}{% else %}[Liste des pièces]{% endif %}"
    )

    doc.add_paragraph()

    # ---- Signature ----
    _add_heading(doc, "SIGNATURE", 2)
    doc.add_paragraph(
        "Fait à {{ ville_redaction }}, le {{ date_redaction }}"
    )
    doc.add_paragraph()
    doc.add_paragraph("{{ nom_redacteur }}")

    # ---- Pied de page : date de génération ----
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "Document généré le {{ date_generation }}"

    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print(f"  ✓  {path}")


# ---------------------------------------------------------------------------
# Point d'entrée
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Création des fichiers d'exemple…\n")

    create_questions_excel("config/questions.xlsx")
    create_template("templates/template.docx", variant="standard")
    create_template("templates/template_mineur.docx", variant="mineur")

    print("\nTerminé ! Vous pouvez maintenant lancer :")
    print("    python main.py")
